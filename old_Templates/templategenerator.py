import re
import os
import tkinter as tk
from tkinter import filedialog, messagebox
from copy import deepcopy
from html import unescape
from docx import Document

def sanitize_tag_name(tag):
    """Convert placeholder text into a valid Jinja variable name."""
    cleaned = unescape(tag).strip()
    cleaned = re.sub(r"<<|>>", "", cleaned)
    cleaned = cleaned.replace("&", " and ")
    cleaned = re.sub(r"[^0-9a-zA-Z]+", "_", cleaned)
    cleaned = re.sub(r"_+", "_", cleaned).strip("_").lower()

    if not cleaned:
        cleaned = "field"
    if cleaned[0].isdigit():
        cleaned = f"field_{cleaned}"

    return cleaned

def set_row_text(row, values):
    for index, cell in enumerate(row.cells):
        cell.text = values[index] if index < len(values) else ""

def insert_control_row(table, row_index, tag_text, before=True):
    reference_row = table.rows[row_index]._tr
    new_row = deepcopy(reference_row)
    insert_index = table._tbl.index(reference_row) if before else table._tbl.index(reference_row) + 1
    table._tbl.insert(insert_index, new_row)
    target_row = table.rows[row_index] if before else table.rows[row_index + 1]
    set_row_text(target_row, [tag_text])

def trim_table_after_row(table, row_index):
    while len(table.rows) > row_index + 1:
        table._tbl.remove(table.rows[row_index + 1]._tr)

def apply_structured_table_loops(doc):
    for table in doc.tables:
        if not table.rows:
            continue
            
        headers = [sanitize_tag_name(cell.text) if cell.text.strip() else '' for cell in table.rows[0].cells]

        if headers == ["", "name", "signature", "position", "date"]:
            trim_table_after_row(table, 1)
            set_row_text(table.rows[1], ["{{ row.role }}", "{{ row.name }}", "{{ row.signature }}", "{{ row.position }}", "{{ row.date }}"])
            insert_control_row(table, 1, "{%tr for row in preparation_table %}", before=True)
            insert_control_row(table, 2, "{%tr endfor %}", before=False)
        elif headers == ["version_no", "new_document", "modified", "date"]:
            trim_table_after_row(table, 1)
            set_row_text(table.rows[1], ["{{ row.version_no }}", "{{ row.new_document }}", "{{ row.modified }}", "{{ row.date }}"])
            insert_control_row(table, 1, "{%tr for row in revision_history_table %}", before=True)
            insert_control_row(table, 2, "{%tr endfor %}", before=False)
        elif headers == ["name", "company", "position", "copy_no", "date"]:
            trim_table_after_row(table, 1)
            set_row_text(table.rows[1], ["{{ row.name }}", "{{ row.company }}", "{{ row.position }}", "{{ row.copy_no }}", "{{ row.date }}"])
            insert_control_row(table, 1, "{%tr for row in distribution_table %}", before=True)
            insert_control_row(table, 2, "{%tr endfor %}", before=False)
        elif headers == ["amendment", "section", "date"]:
            trim_table_after_row(table, 1)
            set_row_text(table.rows[1], ["{{ row.amendment }}", "{{ row.section }}", "{{ row.date }}"])
            insert_control_row(table, 1, "{%tr for row in change_log_table %}", before=True)
            insert_control_row(table, 2, "{%tr endfor %}", before=False)
        elif headers == ["name", "position", "phone", "email"]:
            preview = " ".join(" ".join(cell.text for cell in row.cells) for row in table.rows[:6]).lower()
            loop_name = "badge_contacts_table" if "project manager" in preview and "whs officer" in preview else "tm_consultants_table"
            trim_table_after_row(table, 1)
            set_row_text(table.rows[1], ["{{ row.name }}", "{{ row.position }}", "{{ row.phone }}", "{{ row.email }}"])
            insert_control_row(table, 1, f"{{%tr for row in {loop_name} %}}", before=True)
            insert_control_row(table, 2, "{%tr endfor %}", before=False)
        elif headers == ["department", "position", "phone", "email"]:
            preview = " ".join(" ".join(cell.text for cell in row.cells) for row in table.rows[:6]).lower()
            loop_name = "dtmr_contacts_table" if "dtmr" in preview else "authority_contacts_table"
            trim_table_after_row(table, 1)
            set_row_text(table.rows[1], ["{{ row.department }}", "{{ row.position }}", "{{ row.phone }}", "{{ row.email }}"])
            insert_control_row(table, 1, f"{{%tr for row in {loop_name} %}}", before=True)
            insert_control_row(table, 2, "{%tr endfor %}", before=False)
        elif headers == ["name", "position", "tmr_tmd", "phone", "email"]:
            trim_table_after_row(table, 1)
            set_row_text(table.rows[1], ["{{ row.name }}", "{{ row.position }}", "{{ row.tmr_tmd }}", "{{ row.phone }}", "{{ row.email }}"])
            insert_control_row(table, 1, "{%tr for row in nto_contacts_table %}", before=True)
            insert_control_row(table, 2, "{%tr endfor %}", before=False)
        elif headers == ["name", "position", "phone"]:
            preview = " ".join(" ".join(cell.text for cell in row.cells) for row in table.rows[:6]).lower()
            loop_name = "traffic_control_provider_table" if "site manager" in preview or "site supervisor" in preview else "emergency_contacts_table"
            trim_table_after_row(table, 1)
            set_row_text(table.rows[1], ["{{ row.name }}", "{{ row.position }}", "{{ row.phone }}"])
            insert_control_row(table, 1, f"{{%tr for row in {loop_name} %}}", before=True)
            insert_control_row(table, 2, "{%tr endfor %}", before=False)

def convert_to_jinja_tags(input_file, output_file):
    print(f"Processing: {os.path.basename(input_file)}...")
    doc = Document(input_file)
    pattern = re.compile(r"<<(.*?)>>")

    def replace_text_in_paragraph(paragraph):
        full_text = paragraph.text
        if pattern.search(full_text):
            new_text = pattern.sub(lambda m: "{{ " + sanitize_tag_name(m.group(1)) + " }}", full_text)
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)

    for para in doc.paragraphs:
        replace_text_in_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_in_paragraph(para)

    apply_structured_table_loops(doc)
    doc.save(output_file)
    print(f"Success! Saved to: {output_file}\n")

if __name__ == "__main__":
    # Initialize Tkinter and hide the main window
    root = tk.Tk()
    root.withdraw()

    # Ask user to select files
    print("Please select the Word documents you wish to convert...")
    selected_files = filedialog.askopenfilenames(
        title="Select Word Templates",
        filetypes=[("Word Documents", "*.docx")]
    )

    if not selected_files:
        print("No files selected. Exiting.")
    else:
        # Define output folder
        base_dir = os.path.dirname(selected_files[0])
        output_folder = os.path.abspath(os.path.join(base_dir, "New_Templates"))
        
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

        for input_path in selected_files:
            filename = os.path.basename(input_path)
            output_path = os.path.join(output_folder, filename)
            convert_to_jinja_tags(input_path, output_path)

        messagebox.showinfo("Done", f"Processing complete!\nFiles saved in: {output_folder}")
        print(f"All done! Check the '{output_folder}' folder.")