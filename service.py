import base64
import io
import os
import re
import zipfile
from collections import Counter
from datetime import datetime
from functools import lru_cache

import fitz
from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm
from docxtpl import DocxTemplate, InlineImage

app = Flask(__name__)
app.config['DEBUG'] = os.getenv('FLASK_DEBUG', 'false').lower() == 'true'
CORS(app, origins=['http://127.0.0.1:5000'])

BASE_DIR = os.path.dirname(__file__)
TEMPLATE_DIR = os.path.join(BASE_DIR, 'New_Templates')
HTML_FILE = os.path.join(BASE_DIR, 'index.html')
TEMPLATE_NAME_ALIASES = {
    'Full_TMP.docx': '(FULL) CC00000-TMP-Rev00.docx',
    'Medium_TMP.docx': '(MED) CC00000-TMP-Rev00.docx',
    'Mini_TMP.docx': '(MINI) CC00000-TMP-Rev00.docx',
    'CC00000-TMP-Rev00-Jinja.docx': '(FULL) CC00000-TMP-Rev00.docx',
}

BASE_FORM_FIELDS = {
    'project_name', 'project_location', 'local_government_area', 'client_company',
    'client_contact', 'client_position', 'client_phone', 'client_email',
    'project_description', 'scope_of_work', 'construction_activity', 'start_date',
    'end_date', 'duration', 'pm_name', 'pm_phone', 'pm_email', 'sm_name',
    'sm_phone', 'sm_email', 'se_name', 'se_phone', 'se_email', 'wo_name',
    'wo_phone', 'wo_email', 'traffic_management_notes', 'tgs_reference',
    'tgs_number', 'stages', 'data_source', 'jurisdiction', 'target_year',
    'hourly_database_data', 'hourly_db_data', 'calculated_estimates', 'risk_assessment_summary',
    'cc', 'cc_number', 'revision_number', 'tmr_cert_82_number', 'author_name', 'author_signature', 'author_position',
    'author_date', 'reviewer_name', 'reviewer_op', 'lot_parcel_number', 'development_type', 'number_of_floors',
    'document_preparation', 'revision_history', 'distribution_list', 'change_log', 'badge_contacts',
    'tm_consultants', 'authority_contacts', 'dtmr_contact', 'dtmr_email', 'dtmr_contacts',
    'council_dept', 'council_officer', 'council_phone', 'council_email',
    'nto_name', 'nto_position', 'nto_tmr_tmd', 'nto_tmd_number', 'nto_op_number', 'nto_phone', 'nto_email',
    'emergency_contacts', 'traffic_control_contacts', 'traffic_control_company', 'tc_provider_name', 'tc_provider_phone'
}

AUTO_FIELD_NAMES = {
    'project_tile', 'project_location_e_g_12_example_road_sampletown',
    'street_address_suburb_state_postcode', 'client_contracting_company_name',
    'tc_company', 'scope_of_works', 'enter_construction_stage',
    'enter_construction_methodology_for_the_stage', 'date', 'enter_date',
    'release_date', 'current_year', 'insert_current_year', 'name', 'position',
    'phone', 'email', 'describe_lga', 'describe_suburb', 'road_name',
    'insert_location', 'insert_road', 'insert_ttm', 'insert_active_ttm_measures',
    'detail_ttm_implemented_for_vehicular_traffic', 'yes_no',
    'adopted_not_adopted', 'affected_unaffected', 'are_affected_or_unaffected',
    'be_affected_be_unaffected'
}

DELETE_FIELD_NAMES = {
    'delete_other_reviewers_as_required', 'remove_if_not_provided',
    'remove_the_equipment_not_relevant_to_project',
    'required_for_gccc_remove_if_not_required'
}

IMAGE_FIELD_NAMES = {
    'insert_figure',
    'insert_caption_to_figure_by_right_clicking_the_image'
}

KNOWN_TGS_STAGES = [
    'Gully Pit Inspection',
    'Service Locating',
    'Barrier Installation',
    'Footpath Works',
    'Crossover Construction',
    'Kerb and Channel Works',
    'Excavation / Trenching',
    'Pavement Reinstatement',
    'Line Marking and Signage',
    'Final Traffic Switch',
]


def to_snake_case(value):
    value = re.sub(r'([a-z0-9])([A-Z])', r'\1_\2', value)
    value = re.sub(r'[^0-9a-zA-Z]+', '_', value)
    return value.strip('_').lower()


def template_mtime(path):
    try:
        return os.path.getmtime(path)
    except OSError:
        return 0


@lru_cache(maxsize=16)
def _extract_template_fields(template_path, mtime):
    with zipfile.ZipFile(template_path) as archive:
        xml = archive.read('word/document.xml').decode('utf-8', errors='ignore')
    return sorted(set(re.findall(r'\{\{\s*(.*?)\s*\}\}', xml)))


def extract_template_fields(template_path):
    return _extract_template_fields(template_path, template_mtime(template_path))


def prettify_field_name(field):
    overrides = {
        'cc': 'CC Number',
        'pm_name': 'Project Manager Name',
        'pm_phone': 'Project Manager Phone',
        'pm_email': 'Project Manager Email',
        'sm_name': 'Site Manager Name',
        'sm_phone': 'Site Manager Phone',
        'sm_email': 'Site Manager Email',
        'se_name': 'Site Engineer Name',
        'se_phone': 'Site Engineer Phone',
        'se_email': 'Site Engineer Email',
        'wo_name': 'Works Officer Name',
        'wo_phone': 'Works Officer Phone',
        'wo_email': 'Works Officer Email',
        'tc_company': 'Traffic Control Company',
        'tgs_reference': 'TGS Reference',
        'tgs_number': 'TGS Number',
        'yes_no': 'Yes / No',
        'project_tile': 'Project Title',
    }
    if field in overrides:
        return overrides[field]

    words = field.replace('_', ' ').split()
    upper_words = {'tgs', 'tmp', 'ttm', 'lga', 'aadt', 'lv', 'hv'}
    return ' '.join(word.upper() if word.lower() in upper_words else word.capitalize() for word in words)


def guess_input_type(field):
    boolean_fields = {
        'yes_no', 'adopted_not_adopted', 'affected_unaffected',
        'are_affected_or_unaffected', 'be_affected_be_unaffected'
    }
    long_text_markers = (
        'description', 'scope', 'details', 'reasoning', 'methodology', 'measure',
        'impact', 'terrain', 'route', 'condition', 'ttm', 'activity', 'constraint'
    )

    if field in boolean_fields:
        return 'select'
    if field == 'date' or field.endswith('_date'):
        return 'date'
    if 'email' in field:
        return 'email'
    if 'phone' in field:
        return 'tel'
    if any(marker in field for marker in long_text_markers):
        return 'textarea'
    return 'text'


def get_field_options(field):
    if 'yes_no' in field:
        return ['Yes', 'No']
    if 'adopted' in field:
        return ['Adopted', 'Not Adopted']
    if 'affected' in field:
        return ['Affected', 'Unaffected']
    return []


def guess_field_section(field):
    if any(word in field for word in ('project', 'client', 'company', 'contact', 'address', 'location', 'suburb', 'lga')):
        return 'Project and Client Details'
    if any(word in field for word in ('pm_', 'sm_', 'se_', 'wo_', 'name', 'email', 'phone', 'position')):
        return 'Team Contacts'
    if any(word in field for word in ('stage', 'construction', 'activity', 'methodology', 'duration', 'start', 'end', 'scope')):
        return 'Construction and Staging'
    if any(word in field for word in ('traffic', 'ttm', 'road', 'lane', 'ped', 'speed', 'delay', 'detour', 'aadt', 'queue', 'los', 'hospital', 'school', 'business')):
        return 'Traffic and Impacts'
    if 'tgs' in field or 'figure' in field or 'caption' in field:
        return 'TGS and Figures'
    return 'Additional Parameters'


def classify_template_field(field):
    if field in IMAGE_FIELD_NAMES or field.startswith('insert_figure'):
        category = 'replace'
        help_text = 'This placeholder is replaced by the uploaded TGS PDF pages or their caption.'
    elif field in DELETE_FIELD_NAMES or field.startswith('e_g_') or field.startswith('remove_') or 'delete_' in field:
        category = 'delete'
        help_text = 'This is an instructional or example placeholder and is removed during cleanup.'
    elif field in BASE_FORM_FIELDS or field in AUTO_FIELD_NAMES:
        category = 'auto'
        help_text = 'This field is already covered by the main form or automatic defaults.'
    else:
        category = 'update'
        help_text = 'Enter a value for this template-specific parameter.'

    return {
        'key': field,
        'label': prettify_field_name(field),
        'category': category,
        'inputType': guess_input_type(field),
        'options': get_field_options(field),
        'section': guess_field_section(field),
        'helpText': help_text,
    }


@lru_cache(maxsize=16)
def _build_template_index(template_name, mtime):
    template_path = os.path.join(TEMPLATE_DIR, template_name)
    fields = extract_template_fields(template_path)

    index = {
        'template': template_name,
        'counts': {'update': 0, 'auto': 0, 'replace': 0, 'delete': 0},
        'fields': {'update': [], 'auto': [], 'replace': [], 'delete': []},
        'sections': [],
    }

    section_map = {}
    for field in fields:
        item = classify_template_field(field)
        index['fields'][item['category']].append(item)
        index['counts'][item['category']] += 1

        if item['category'] == 'update':
            section_map.setdefault(item['section'], []).append(item)

    index['sections'] = [
        {'name': name, 'fields': sorted(items, key=lambda entry: entry['label'])}
        for name, items in sorted(section_map.items())
    ]

    for category in index['fields']:
        index['fields'][category] = sorted(index['fields'][category], key=lambda entry: entry['label'])

    return index


def build_template_index(template_name):
    template_path = os.path.join(TEMPLATE_DIR, template_name)
    return _build_template_index(template_name, template_mtime(template_path))


def iter_paragraphs(parent):
    for paragraph in getattr(parent, 'paragraphs', []):
        yield paragraph
    for table in getattr(parent, 'tables', []):
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def delete_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def normalize_context_value(value):
    if value is None:
        return ''
    if isinstance(value, list):
        return ', '.join(str(item) for item in value if item)
    return str(value)


def resolve_template_name(requested_name):
    requested_name = normalize_context_value(requested_name).strip()
    if requested_name:
        alias_name = TEMPLATE_NAME_ALIASES.get(requested_name, requested_name)
        alias_path = os.path.join(TEMPLATE_DIR, alias_name)
        if os.path.exists(alias_path):
            return alias_name

    for fallback_name in TEMPLATE_NAME_ALIASES.values():
        fallback_path = os.path.join(TEMPLATE_DIR, fallback_name)
        if os.path.exists(fallback_path):
            return fallback_name

    available_templates = sorted(
        filename for filename in os.listdir(TEMPLATE_DIR)
        if filename.lower().endswith('.docx')
    )
    return available_templates[0] if available_templates else requested_name


def fallback_for_key(key, context):
    alias_map = {
        'project_title': 'project_name',
        'project_tile': 'project_name',
        'tc_company': 'tc_company',
        'traffic_control_company': 'tc_company',
        'suburb': 'project_location',
        'road': 'project_location',
        'road_name': 'project_location',
        'local_government_area': 'local_government_area',
        'enter_construction_stage': 'enter_construction_stage',
        'enter_construction_methodology_for_the_stage': 'construction_activity',
    }

    if key in alias_map:
        return normalize_context_value(context.get(alias_map[key], ''))
    if 'company' in key:
        if 'tc' in key or 'traffic_control' in key:
            return normalize_context_value(context.get('tc_company', ''))
        return normalize_context_value(context.get('client_company', ''))
    if 'project' in key and ('name' in key or 'title' in key or 'tile' in key):
        return normalize_context_value(context.get('project_name', ''))
    if 'location' in key or 'road' in key or 'suburb' in key or 'lga' in key:
        return normalize_context_value(context.get('project_location', ''))
    if 'stage' in key or 'activity' in key or 'methodology' in key:
        return normalize_context_value(context.get('enter_construction_stage') or context.get('construction_activity', ''))
    if 'email' in key:
        return normalize_context_value(context.get('email', ''))
    if 'phone' in key:
        return normalize_context_value(context.get('phone', ''))
    if key.endswith('_name') or key == 'name':
        return normalize_context_value(context.get('name', ''))
    return ''


def replace_literal_placeholders(text, context):
    def replace_match(match):
        raw_key = to_snake_case(match.group(1))
        value = normalize_context_value(context.get(raw_key, ''))
        if not value:
            value = fallback_for_key(raw_key, context)
        return value

    updated_text = re.sub(r'<<\s*(.*?)\s*>>', replace_match, text, flags=re.IGNORECASE)

    document_reference = normalize_context_value(context.get('document_reference', ''))
    if document_reference:
        updated_text = re.sub(r'CC\d{5}-TMP-Rev\s*\d+', document_reference, updated_text, flags=re.IGNORECASE)

    updated_text = updated_text.replace('<<', '').replace('>>', '')
    return updated_text


def enable_update_fields(document):
    settings_element = document.settings.element
    update_fields = settings_element.find(qn('w:updateFields'))
    if update_fields is None:
        update_fields = OxmlElement('w:updateFields')
        settings_element.append(update_fields)
    update_fields.set(qn('w:val'), 'true')


def format_date_value(value):
    if not value:
        return ''
    value = str(value).strip()
    for fmt in ('%Y-%m-%d', '%d/%m/%Y', '%d-%m-%Y'):
        try:
            return datetime.strptime(value, fmt).strftime('%d/%m/%Y')
        except ValueError:
            continue
    return value


def clean_scanned_line(value):
    text = normalize_context_value(value)
    text = text.replace('\ufb01', 'fi').replace('\ufb02', 'fl')
    text = re.sub(r'[\x00-\x1f\x7f-\x9f]+', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip(' -–|:\t')


def is_address_like(text):
    cleaned = clean_scanned_line(text)
    if not cleaned:
        return False
    return bool(re.search(r'\b\d+\s+.*\b(?:RD|ROAD|ST|STREET|AVE|AVENUE|DR|DRIVE|HWY|HIGHWAY|CT|COURT|PDE|PARADE|CRES|CRESCENT|PL|PLACE|BLVD|BOULEVARD)\b.*\b(?:QLD|NSW|VIC|SA|WA|TAS|NT|ACT)\b', cleaned, re.IGNORECASE))


def infer_local_government_area(text):
    cleaned = clean_scanned_line(text).upper()
    lga_hints = {
        'BRISBANE': 'Brisbane City Council',
        'GOLD COAST': 'Gold Coast City Council',
        'LOGAN': 'Logan City Council',
        'IPSWICH': 'Ipswich City Council',
        'MORETON BAY': 'Moreton Bay City Council',
        'SUNSHINE COAST': 'Sunshine Coast Council',
        'GYMPIE': 'Gympie Regional Council',
        'REDLAND': 'Redland City Council',
    }
    for hint, lga in lga_hints.items():
        if hint in cleaned:
            return lga
    return ''


def is_valid_stage_candidate(text):
    cleaned = clean_scanned_line(text)
    if not cleaned or len(cleaned) < 4 or len(cleaned) > 80:
        return False

    upper = cleaned.upper()
    if upper in {stage.upper() for stage in KNOWN_TGS_STAGES}:
        return True

    if re.search(r'\.{4,}', cleaned) or re.fullmatch(r'\d+(?:\s*OF\s*\d+)?', upper):
        return False
    if upper.startswith(('PAGE ', 'FIGURE ', 'TABLE ', 'NOTE', 'NOTES', 'DESCRIPTION', 'CONTACT', 'ROAD SPEED', 'BETWEEN ')):
        return False
    if upper in {'DESCRIPTION', 'NOTES', 'CONTACT', 'RISK ASSESSMENT', 'PLAN', 'ROAD', 'RD'}:
        return False
    if is_address_like(cleaned):
        return False

    return any(keyword in upper for keyword in (
        'WORKS', 'INSTALLATION', 'INSPECTION', 'LOCATING', 'CONSTRUCTION',
        'TRENCH', 'REINSTATEMENT', 'MARKING', 'SWITCH', 'EXCAVATION'
    ))


def is_valid_methodology_candidate(text):
    cleaned = clean_scanned_line(text)
    if not cleaned or len(cleaned) < 8 or len(cleaned) > 160:
        return False

    upper = cleaned.upper()
    if re.search(r'\.{4,}', cleaned):
        return False
    if upper.startswith(('FIGURE ', 'TABLE ', 'PAGE ', 'DESCRIPTION', 'NOTES', 'CONTACT')):
        return False
    if 'ROAD SPEED' in upper or upper.startswith('BETWEEN ') or is_address_like(cleaned):
        return False

    return any(keyword in upper for keyword in (
        'PEDESTRIAN', 'LANE', 'CLOSURE', 'DETOUR', 'STOP/SLOW', 'STOP SLOW',
        'WORKS', 'TRAFFIC MANAGEMENT', 'TRAFFIC CONTROL', 'FOOTPATH'
    ))


def best_candidate(values):
    cleaned_values = [clean_scanned_line(value) for value in values if clean_scanned_line(value)]
    if not cleaned_values:
        return ''
    counts = Counter(cleaned_values)
    return sorted(counts.items(), key=lambda item: (-item[1], -len(item[0])))[0][0]


def extract_tgs_metadata_from_text(page_text):
    lines = [clean_scanned_line(line) for line in str(page_text or '').splitlines() if clean_scanned_line(line)]
    metadata = {
        'suite': '',
        'stage_name': '',
        'methodology': '',
        'site_location': '',
        'cross_streets': '',
        'road_speed': '',
        'sheet_reference': '',
        'project_name': '',
        'client_company': '',
        'traffic_control_company': '',
        'local_government_area': '',
    }

    if not lines:
        return metadata

    sheet_pattern = re.compile(r'CC\d{5}-S\d+-?\d+[A-C]?', re.IGNORECASE)
    company_pattern = re.compile(r'\b(?:PTY\s+LTD|CIVIL|CONSTRUCTION|CONTRACTORS?|BUILDERS?|DEVELOPMENTS?|TRAFFIC)\b', re.IGNORECASE)
    project_pattern = re.compile(r'\b(?:PROJECT|SITE|LOCATION|ADDRESS)\s*[:\-]\s*(.+)', re.IGNORECASE)

    for line in lines[:40]:
        upper = line.upper()

        if not metadata['sheet_reference']:
            match = sheet_pattern.search(upper)
            if match:
                metadata['sheet_reference'] = match.group(0).upper()

        if not metadata['road_speed'] and 'ROAD SPEED' in upper:
            metadata['road_speed'] = line

        if not metadata['cross_streets'] and upper.startswith('BETWEEN '):
            metadata['cross_streets'] = line

        if not metadata['suite'] and any(word in upper for word in ('CIVIL', 'CONSTRUCTION', 'DRAINAGE', 'ROADWORK', 'TRAFFIC MANAGEMENT')):
            if len(line) <= 80 and not re.search(r'\.{4,}', line) and 'ROAD SPEED' not in upper:
                metadata['suite'] = line.replace(' - ', ' ').strip()

        if not metadata['stage_name']:
            for stage in KNOWN_TGS_STAGES:
                if stage.upper() in upper:
                    metadata['stage_name'] = stage
                    break
            if not metadata['stage_name'] and is_valid_stage_candidate(line):
                metadata['stage_name'] = line

        if not metadata['methodology'] and is_valid_methodology_candidate(line):
            metadata['methodology'] = line if line.startswith('(') else f'({line})'

        if not metadata['site_location'] and is_address_like(line):
            metadata['site_location'] = line
            metadata['local_government_area'] = metadata['local_government_area'] or infer_local_government_area(line)
            if not metadata['project_name']:
                project_parts = re.split(r'\s+[–-]\s+', line, maxsplit=1)
                if len(project_parts) == 2 and not is_address_like(project_parts[0]):
                    metadata['project_name'] = clean_scanned_line(project_parts[0])
                    metadata['site_location'] = clean_scanned_line(project_parts[1])

        if not metadata['project_name']:
            project_match = project_pattern.search(line)
            if project_match:
                candidate = clean_scanned_line(project_match.group(1))
                if candidate and not is_address_like(candidate):
                    metadata['project_name'] = candidate

        if not metadata['client_company'] and company_pattern.search(line) and 'CROMPTON CONCEPTS' not in upper:
            if len(line) <= 90 and not upper.startswith(('FIGURE ', 'PAGE ', 'TABLE ')):
                metadata['client_company'] = line

        if not metadata['traffic_control_company'] and 'TRAFFIC' in upper and 'CROMPTON CONCEPTS' not in upper:
            if company_pattern.search(line) and len(line) <= 90:
                metadata['traffic_control_company'] = line

    metadata['local_government_area'] = metadata['local_government_area'] or infer_local_government_area(metadata['site_location'])

    if not metadata['methodology'] and metadata['stage_name']:
        metadata['methodology'] = f'({metadata["stage_name"]} traffic management works)'

    return metadata


def build_stage_option_analysis(stages, suite='', methodology=''):
    unique_stages = []
    for stage in stages or []:
        stage_text = normalize_context_value(stage).strip()
        if stage_text and stage_text not in unique_stages:
            unique_stages.append(stage_text)

    if not unique_stages:
        unique_stages = ['General traffic management stage']

    work_suite = normalize_context_value(suite).strip() or 'Civil Construction'
    control_method = normalize_context_value(methodology).strip() or 'the approved TGS arrangement'

    return '\n\n'.join(
        f'Option Analysis - {stage}: For the {work_suite} suite, the preferred methodology implemented is {control_method}. This option maintains safe separation around the work area while minimising impacts on vehicles, pedestrians, and property access.'
        for stage in unique_stages
    )


def parse_numeric_volume(value):
    try:
        if value is None or str(value).strip() == '':
            return None
        return float(str(value).replace(',', '').strip())
    except (TypeError, ValueError):
        return None


# --- Traffic Impact Assessment (TIA) Engine ---
def calculate_traffic_impact(jurisdiction, hourly_database_data, calculated_estimates, target_year, base_volume=12199):
    """
    Generates the AADT, VCR, and Queue Length estimations matching Tables 17-19.
    """
    active_volume_data = parse_numeric_volume(hourly_database_data)
    if active_volume_data is None:
        active_volume_data = parse_numeric_volume(calculated_estimates)
    if active_volume_data is None:
        active_volume_data = float(base_volume)

    current_year = datetime.now().year
    try:
        target_year_int = int(float(target_year))
    except (TypeError, ValueError):
        target_year_int = current_year

    growth_years = target_year_int - current_year if target_year_int > current_year else 0
    projected_target_volume = int(round(active_volume_data * ((1 + 0.025) ** growth_years)))

    if jurisdiction == 'Gold Coast':
        capacity = 1500
        vcr = projected_target_volume / capacity
        queue_length = 150
        return {
            'aadt_target_year': projected_target_volume,
            'vcr_score': round(vcr, 2),
            'queue_length_m': queue_length,
            'los_rating': 'D' if vcr > 0.85 else 'C'
        }

    capacity = 1670
    vcr = projected_target_volume / capacity
    queue_length = 130
    return {
        'aadt_target_year': projected_target_volume,
        'vcr_score': round(vcr, 2),
        'queue_length_m': queue_length,
        'los_rating': 'C' if vcr <= 0.9 else 'D'
    }


# --- Option Analysis Generator ---
def generate_options_analysis(stages):
    """
    Generates the structured Option Analysis tables for Vehicular, Vulnerable Users, Bus Stops, and Property (Tables 13-16).
    """
    if isinstance(stages, str):
        stage_list = [stages.strip()] if stages.strip() else []
    else:
        stage_list = [normalize_context_value(stage).strip() for stage in (stages or []) if normalize_context_value(stage).strip()]

    stage_text = ', '.join(stage_list) if stage_list else 'general site works'
    vehicular = [
        {'option': 'Side-track', 'features': 'Would allow closure of entire carriageway...', 'comment': 'Not practical due to nature of works'},
        {'option': 'Contra Flow', 'features': 'Traffic through the worksite', 'comment': 'Reasonable detour route unavailable with maximum 5 min delay'},
        {'option': 'Hold and Release', 'features': 'Temporarily holding traffic to allow construction vehicles...', 'comment': f'Adopted for {stage_text.lower()}'}
    ]
    pedestrian = [
        {'option': 'Close footpath', 'reasoning': 'Allows elimination of the risk of trips and falls...', 'comment': 'Footpath will be closed. Pedestrian detour implemented.'}
    ]
    return {
        'vehicular_options': vehicular,
        'vulnerable_users': pedestrian,
        'vehicular': vehicular,
        'pedestrian': pedestrian,
    }


# --- Desktop Risk Assessment Generator ---
def generate_desktop_risk_assessment():
    """
    Generates the Risk Assessment Hazard Matrix (Appendix C).
    """
    return [
        {
            'hazard': 'Pedestrians, cyclists, people with disabilities...',
            'potential_risk': 'Unable to pass safely past the site',
            'initial_risk': 'Very High',
            'control_measure': 'Safe access maintained at all times around the site.',
            'residual_risk': 'Medium'
        },
        {
            'hazard': 'Traffic queues and delays',
            'potential_risk': 'Unacceptably long delays to road users',
            'initial_risk': 'Low',
            'control_measure': 'No queuing of traffic required outside allowed thresholds.',
            'residual_risk': 'Low'
        },
        {
            'hazard': 'Complete closure of turning lanes',
            'potential_risk': 'Removal of option for road users',
            'initial_risk': 'High',
            'control_measure': 'Left turning lane closed for 10 days only. Left turns permitted from traffic lane.',
            'residual_risk': 'Medium'
        }
    ]


def calculate_tia(jurisdiction, base_volume=None, target_year=None):
    tia_results = calculate_traffic_impact(jurisdiction, base_volume, base_volume, target_year)
    return {
        'aadt': tia_results.get('aadt_target_year'),
        'vcr': tia_results.get('vcr_score'),
        'los': tia_results.get('los_rating'),
        'queue_length': tia_results.get('queue_length_m'),
    }


def generate_risk_matrix():
    return generate_desktop_risk_assessment()


def generate_ctmp_report(form_data):
    resolved_template_name = resolve_template_name(form_data.get('templateType', 'CC00000-TMP-Rev00-Jinja.docx'))
    template_path = os.path.join(TEMPLATE_DIR, resolved_template_name)
    doc = DocxTemplate(template_path)

    stages = form_data.getlist('stages') if hasattr(form_data, 'getlist') else []
    if not stages:
        construction_stage = normalize_context_value(form_data.get('construction_stage', '')).strip()
        if construction_stage:
            stages = [construction_stage]

    template_fields = extract_template_fields(template_path)
    context = build_context(form_data, template_fields, stages, doc, uploaded_images=[])

    jurisdiction = form_data.get('jurisdiction', 'TMR')
    try:
        target_year = int(form_data.get('target_year', form_data.get('targetYear', datetime.now().year + 2)))
    except (TypeError, ValueError):
        target_year = datetime.now().year + 2

    tia_results = calculate_tia(
        jurisdiction=jurisdiction,
        base_volume=form_data.get('aadt_volume', form_data.get('hourlyDbData') or form_data.get('hourly_database_data')),
        target_year=target_year
    )
    options_analysis = generate_options_analysis(stages or form_data.get('construction_stage', ''))
    risk_matrix = generate_risk_matrix()

    context.update({
        'project_name': form_data.get('project_name', context.get('project_name')),
        'scope_of_works': form_data.get('scope_of_works', context.get('scope_of_works')),
        'project_location': form_data.get('project_location', context.get('project_location')),
        'date': datetime.now().strftime('%d.%m.%Y'),
        'author_name': form_data.get('author_name', context.get('author_name', 'Sanju Bhandari')),
        'client_company': form_data.get('client_company', context.get('client_company')),
        'tia_target_year': target_year,
        'tia_aadt': tia_results.get('aadt'),
        'tia_vcr': tia_results.get('vcr'),
        'tia_los': tia_results.get('los'),
        'tia_queue_length': tia_results.get('queue_length'),
        'vehicular_options': options_analysis.get('vehicular'),
        'pedestrian_options': options_analysis.get('pedestrian'),
        'risk_matrix': risk_matrix,
        'risk_assessment_table': risk_matrix,
    })

    doc.render(context)
    safe_project_name = normalize_context_value(context.get('project_name', 'Project')).replace(' ', '_')
    output_filename = os.path.join(BASE_DIR, f'CTMP_{safe_project_name}.docx')
    doc.save(output_filename)
    return output_filename


def summarize_tgs_analysis(uploaded_images):
    stage_candidates = []
    suite_candidates = []
    methodology_candidates = []
    site_candidates = []
    cross_street_candidates = []
    speed_candidates = []
    sheet_candidates = []
    project_candidates = []
    client_candidates = []
    traffic_company_candidates = []
    lga_candidates = []

    for item in uploaded_images or []:
        stage_name = clean_scanned_line(item.get('stage_name', ''))
        if is_valid_stage_candidate(stage_name):
            stage_candidates.append(stage_name)

        suite_candidates.append(item.get('suite', ''))

        methodology = clean_scanned_line(item.get('methodology', ''))
        if is_valid_methodology_candidate(methodology):
            methodology_candidates.append(methodology)

        site_candidates.append(item.get('site_location', ''))
        cross_street_candidates.append(item.get('cross_streets', ''))
        speed_candidates.append(item.get('road_speed', ''))
        sheet_candidates.append(item.get('sheet_reference', ''))
        project_candidates.append(item.get('project_name', ''))
        client_candidates.append(item.get('client_company', ''))
        traffic_company_candidates.append(item.get('traffic_control_company', ''))
        lga_candidates.append(item.get('local_government_area', ''))

    recognized_stages = [stage for stage in KNOWN_TGS_STAGES if any(stage.lower() == candidate.lower() for candidate in stage_candidates)]
    additional_stages = []
    for candidate in stage_candidates:
        if candidate not in recognized_stages and candidate not in additional_stages:
            additional_stages.append(candidate)
    stages = recognized_stages + additional_stages[:4]

    suite = best_candidate(suite_candidates) or 'Construction'
    methodology = best_candidate(methodology_candidates)
    site_location = best_candidate(site_candidates)
    cross_streets = best_candidate(cross_street_candidates)
    road_speed = best_candidate(speed_candidates)
    sheet_reference = best_candidate(sheet_candidates)
    project_name = best_candidate(project_candidates)
    client_company = best_candidate(client_candidates)
    traffic_control_company = best_candidate(traffic_company_candidates)
    local_government_area = best_candidate(lga_candidates) or infer_local_government_area(site_location)

    if not methodology and stages:
        methodology = f'({stages[0]} traffic management works)'

    option_analysis = build_stage_option_analysis(stages, suite, methodology)

    return {
        'suite': suite,
        'stages': stages,
        'methodology': methodology,
        'siteLocation': site_location,
        'crossStreets': cross_streets,
        'roadSpeed': road_speed,
        'sheetReference': sheet_reference,
        'projectName': project_name,
        'clientCompany': client_company,
        'trafficControlCompany': traffic_control_company,
        'localGovernmentArea': local_government_area,
        'optionAnalysis': option_analysis,
    }


def build_document_reference(cc_value, revision_value):
    cc_value = normalize_context_value(cc_value).strip()
    revision_value = normalize_context_value(revision_value).strip()

    if not cc_value and not revision_value:
        return ''

    if re.search(r'-TMP-REV\s*\d+', cc_value, flags=re.IGNORECASE):
        return cc_value

    if cc_value and revision_value:
        if re.match(r'^CC\d{5}$', cc_value, flags=re.IGNORECASE):
            return f'{cc_value}-TMP-Rev {revision_value}'
        if 'TMP' not in cc_value.upper() and 'REV' not in cc_value.upper():
            return f'{cc_value}-TMP-Rev {revision_value}'
        if 'REV' not in cc_value.upper():
            return f'{cc_value}-Rev {revision_value}'

    return cc_value or (f'TMP-Rev {revision_value}' if revision_value else '')


def initials_from_name(name):
    parts = [part for part in str(name).split() if part]
    if not parts:
        return ''
    if len(parts) == 1:
        return f'{parts[0][0].upper()}.'
    return f'{parts[0][0].upper()}. {parts[-1]}'


def parse_structured_rows(raw_text, expected_columns, default_rows=None):
    rows = []
    for line in str(raw_text or '').splitlines():
        if not line.strip():
            continue
        parts = [part.strip() for part in line.split('|')]
        while len(parts) < expected_columns:
            parts.append('')
        rows.append(parts[:expected_columns])
    return rows if rows else list(default_rows or [])


def rows_to_dicts(rows, keys):
    result = []
    for row in rows:
        item = {}
        for index, key in enumerate(keys):
            item[key] = normalize_context_value(row[index]) if index < len(row) else ''
        result.append(item)
    return result


def build_loop_context(structured_tables):
    preparation_rows = rows_to_dicts(structured_tables['document_preparation'], ['role', 'name', 'signature', 'position', 'date'])

    revision_rows = []
    for row in structured_tables['revision_history']:
        version = normalize_context_value(row[0]) if len(row) > 0 else ''
        document = normalize_context_value(row[1]) if len(row) > 1 else ''
        modified = normalize_context_value(row[2]) if len(row) > 2 else ''
        date = normalize_context_value(row[3]) if len(row) > 3 else ''
        revision_rows.append({
            'version_no': version,
            'version': version,
            'new_document': document,
            'document': document,
            'revision': modified,
            'modified': modified,
            'date': date,
        })

    distribution_rows = rows_to_dicts(structured_tables['distribution'], ['name', 'company', 'position', 'copy_no', 'date'])
    change_log_rows = rows_to_dicts(structured_tables['change_log'], ['amendment', 'section', 'date'])
    badge_rows = rows_to_dicts(structured_tables['badge_contacts'], ['name', 'position', 'phone', 'email'])
    consultant_rows = rows_to_dicts(structured_tables['tm_consultants'], ['name', 'position', 'phone', 'email'])
    authority_rows = rows_to_dicts(structured_tables['authority_contacts'], ['department', 'position', 'phone', 'email'])
    dtmr_rows = rows_to_dicts(structured_tables['dtmr_contacts'], ['department', 'position', 'phone', 'email'])

    nto_rows = []
    for row in structured_tables['nto_contacts']:
        tmd_number = normalize_context_value(row[2]) if len(row) > 2 else ''
        nto_rows.append({
            'name': normalize_context_value(row[0]) if len(row) > 0 else '',
            'position': normalize_context_value(row[1]) if len(row) > 1 else '',
            'tmr_tmd': tmd_number,
            'tmd_number': tmd_number,
            'phone': normalize_context_value(row[3]) if len(row) > 3 else '',
            'email': normalize_context_value(row[4]) if len(row) > 4 else '',
        })

    emergency_rows = rows_to_dicts(structured_tables['emergency_contacts'], ['name', 'position', 'phone'])
    traffic_control_rows = rows_to_dicts(structured_tables['traffic_control_contacts'], ['name', 'position', 'phone'])

    return {
        'preparation_table': preparation_rows,
        'document_preparation_table': preparation_rows,
        'approval_table': preparation_rows,
        'reviewers_table': preparation_rows,
        'revision_table': revision_rows,
        'revision_history_table': revision_rows,
        'distribution_table': distribution_rows,
        'change_log_table': change_log_rows,
        'badge_contacts_table': badge_rows,
        'contacts_badge': badge_rows,
        'tm_consultants_table': consultant_rows,
        'authority_contacts_table': authority_rows,
        'contacts_council': authority_rows,
        'dtmr_contacts_table': dtmr_rows,
        'nto_contacts_table': nto_rows,
        'contacts_nto': nto_rows,
        'emergency_contacts_table': emergency_rows,
        'traffic_control_provider_table': traffic_control_rows,
        'traffic_control_contacts_table': traffic_control_rows,
        'contacts_tc_provider': traffic_control_rows,
    }


def write_rows_to_table(table, rows, start_row=1):
    while len(table.rows) < start_row + len(rows):
        table.add_row()

    for row_index, values in enumerate(rows, start=start_row):
        row = table.rows[row_index]
        for cell_index, cell in enumerate(row.cells):
            cell.text = normalize_context_value(values[cell_index]) if cell_index < len(values) else ''

    for row_index in range(start_row + len(rows), len(table.rows)):
        for cell in table.rows[row_index].cells:
            cell.text = ''


def table_headers(table):
    return [to_snake_case(' '.join(cell.text.split())) for cell in table.rows[0].cells]


def populate_structured_tables(document, context):
    structured = context.get('_structured_tables', {})

    for table in document.tables:
        headers = table_headers(table)
        table_preview = ' '.join(
            ' '.join(' '.join(cell.text.split()) for cell in row.cells)
            for row in table.rows[:6]
        ).lower()

        if headers == ['', 'name', 'signature', 'position', 'date']:
            write_rows_to_table(table, structured.get('document_preparation', []), start_row=1)
        elif headers == ['version_no', 'new_document', 'modified', 'date']:
            write_rows_to_table(table, structured.get('revision_history', []), start_row=1)
        elif headers == ['name', 'company', 'position', 'copy_no', 'date']:
            write_rows_to_table(table, structured.get('distribution', []), start_row=1)
        elif headers == ['amendment', 'section', 'date']:
            write_rows_to_table(table, structured.get('change_log', []), start_row=1)
        elif headers == ['name', 'position', 'phone', 'email']:
            if 'project manager' in table_preview and 'whs officer' in table_preview:
                write_rows_to_table(table, structured.get('badge_contacts', []), start_row=1)
            elif 'traffic technician' in table_preview or 'cromptonconcepts' in table_preview or 'tmd' in table_preview:
                write_rows_to_table(table, structured.get('tm_consultants', []), start_row=1)
        elif headers == ['department', 'position', 'phone', 'email']:
            if 'dtmr' in table_preview:
                write_rows_to_table(table, structured.get('dtmr_contacts', []), start_row=1)
            else:
                write_rows_to_table(table, structured.get('authority_contacts', []), start_row=1)
        elif headers == ['name', 'position', 'tmr_tmd', 'phone', 'email']:
            write_rows_to_table(table, structured.get('nto_contacts', []), start_row=1)
        elif headers == ['name', 'position', 'phone']:
            if 'site manager' in table_preview or 'site supervisor' in table_preview:
                write_rows_to_table(table, structured.get('traffic_control_contacts', []), start_row=1)
            else:
                write_rows_to_table(table, structured.get('emergency_contacts', []), start_row=1)
        elif len(headers) >= 2 and headers[1] == 'details':
            for row in table.rows[1:]:
                label = to_snake_case(row.cells[0].text)
                if label == 'title':
                    row.cells[1].text = normalize_context_value(context.get('project_name', ''))
                elif label == 'description':
                    row.cells[1].text = normalize_context_value(context.get('brief_description_of_the_project_including_nature_goals_and_any_major_components', ''))
                elif label == 'development_footprint':
                    row.cells[1].text = normalize_context_value(context.get('number_of_floors_e_g_3_storeys', ''))
                elif label == 'type':
                    row.cells[1].text = normalize_context_value(context.get('type_of_development_e_g_residential_commercial_mixed_use', ''))
                elif label == 'address':
                    row.cells[1].text = normalize_context_value(context.get('street_address_suburb_state_postcode', ''))
                elif label == 'project_location':
                    row.cells[1].text = normalize_context_value(context.get('project_location', ''))
                elif label == 'lot_parcel_number':
                    row.cells[1].text = normalize_context_value(context.get('insert_lot_and_parcel_number', ''))


def build_context(form, template_fields, stages, doc, uploaded_images):
    now = datetime.now()
    today = now.strftime('%d/%m/%Y')
    current_year = now.strftime('%Y')

    normalized_form = {}
    for key in form.keys():
        values = [value.strip() for value in form.getlist(key) if value and value.strip()]
        if not values:
            continue
        normalized_key = to_snake_case(key)
        normalized_value = ', '.join(values) if len(values) > 1 else values[0]
        if 'date' in normalized_key:
            normalized_value = format_date_value(normalized_value)
        normalized_form[normalized_key] = normalized_value

    if 'tmr_cert82_number' in normalized_form and 'tmr_cert_82_number' not in normalized_form:
        normalized_form['tmr_cert_82_number'] = normalized_form['tmr_cert82_number']

    project_name = normalized_form.get('project_name', 'Project')
    project_location = normalized_form.get('project_location', '')
    client_company = normalized_form.get('client_company', '')
    tc_provider_name = normalized_form.get('tc_provider_name', '')
    tc_provider_phone = normalized_form.get('tc_provider_phone', '')
    traffic_control_company = normalized_form.get('traffic_control_company', normalized_form.get('tc_company', tc_provider_name or client_company))

    client_name = normalized_form.get('client_contact', '')
    client_position = normalized_form.get('client_position', '')
    client_phone = normalized_form.get('client_phone', '')
    client_email = normalized_form.get('client_email', '')
    reviewer_name = normalized_form.get('reviewer_name', 'Aaron Anthony')
    reviewer_op = normalized_form.get('reviewer_op', 'OP611')
    council_dept = normalized_form.get('council_dept', 'Traffic and Transport')
    council_officer = normalized_form.get('council_officer', 'Assessing Officer')
    council_phone = normalized_form.get('council_phone', '')
    council_email = normalized_form.get('council_email', '')

    pm_name = normalized_form.get('pm_name', client_name)
    pm_position = 'Project Manager'
    pm_phone = normalized_form.get('pm_phone', client_phone)
    pm_email = normalized_form.get('pm_email', client_email)

    sm_name = normalized_form.get('sm_name', '')
    sm_position = 'Site Manager'
    sm_phone = normalized_form.get('sm_phone', '')
    sm_email = normalized_form.get('sm_email', '')

    se_name = normalized_form.get('se_name', '')
    se_position = 'Site Engineer'
    se_phone = normalized_form.get('se_phone', '')
    se_email = normalized_form.get('se_email', '')

    wo_name = normalized_form.get('wo_name', '')
    wo_position = 'Site Supervisor'
    wo_phone = normalized_form.get('wo_phone', '')
    wo_email = normalized_form.get('wo_email', '')

    nto_name = normalized_form.get('nto_name', client_name)
    nto_position = normalized_form.get('nto_position', 'TMD' if normalized_form.get('nto_name') else client_position)
    nto_tmd_number = normalized_form.get('nto_tmd_number', '')
    nto_op_number = normalized_form.get('nto_op_number', '')
    nto_tmr_tmd = normalized_form.get('nto_tmr_tmd', nto_op_number or nto_tmd_number)
    nto_phone = normalized_form.get('nto_phone', '')
    nto_email = normalized_form.get('nto_email', '')

    generic_name = client_name or project_name
    generic_position = client_position
    generic_phone = client_phone
    generic_email = client_email

    tgs_metadata = {
        'suite': '',
        'stage_name': '',
        'methodology': '',
        'site_location': '',
        'cross_streets': '',
        'road_speed': '',
        'sheet_reference': '',
        'project_name': '',
        'client_company': '',
        'traffic_control_company': '',
        'local_government_area': '',
    }
    for image_info in uploaded_images or []:
        for key in tgs_metadata:
            if not tgs_metadata[key] and image_info.get(key):
                tgs_metadata[key] = normalize_context_value(image_info.get(key)).strip()

    tgs_suite = tgs_metadata.get('suite', '')
    tgs_stage_name = tgs_metadata.get('stage_name', '')
    tgs_methodology = tgs_metadata.get('methodology', '')
    tgs_site_location = tgs_metadata.get('site_location', '')
    tgs_cross_streets = tgs_metadata.get('cross_streets', '')
    tgs_road_speed = tgs_metadata.get('road_speed', '')
    tgs_sheet_reference = tgs_metadata.get('sheet_reference', '')
    tgs_project_name = tgs_metadata.get('project_name', '')
    tgs_client_company = tgs_metadata.get('client_company', '')
    tgs_traffic_control_company = tgs_metadata.get('traffic_control_company', '')
    tgs_local_government_area = tgs_metadata.get('local_government_area', '')

    if not project_name and tgs_project_name:
        project_name = tgs_project_name
    if not project_location and tgs_site_location:
        project_location = tgs_site_location
    if not client_company and tgs_client_company:
        client_company = tgs_client_company
    if not traffic_control_company and tgs_traffic_control_company:
        traffic_control_company = tgs_traffic_control_company

    if tgs_stage_name and tgs_stage_name not in stages:
        stages = [*stages, tgs_stage_name]

    project_description = normalized_form.get('project_description', '')
    scope_of_work = normalized_form.get('scope_of_work', project_description)
    construction_activity = normalized_form.get('construction_activity', tgs_methodology or scope_of_work or ', '.join(stages))
    stage_text = ', '.join(stages) if stages else (tgs_stage_name or construction_activity or 'General construction works')
    start_date = normalized_form.get('start_date', today)
    end_date = normalized_form.get('end_date', today)
    duration = normalized_form.get('duration', '')
    local_government_area = normalized_form.get('local_government_area', tgs_local_government_area)
    jurisdiction = normalized_form.get('jurisdiction', 'TMR')
    target_year = normalized_form.get('target_year', current_year)
    hourly_database_data = parse_numeric_volume(
        normalized_form.get('hourly_database_data', normalized_form.get('hourly_db_data'))
    )
    calculated_estimates = parse_numeric_volume(normalized_form.get('calculated_estimates'))
    tia_metrics = calculate_traffic_impact(
        jurisdiction=jurisdiction,
        hourly_database_data=hourly_database_data,
        calculated_estimates=calculated_estimates,
        target_year=target_year,
    )
    structured_options = generate_options_analysis(stages)
    desktop_risk_matrix = generate_desktop_risk_assessment()
    risk_assessment_summary = normalized_form.get('risk_assessment_summary', '')
    if not risk_assessment_summary:
        risk_assessment_summary = '\n'.join(
            f"{item['hazard']}: {item['control_measure']} (Residual risk: {item['residual_risk']})"
            for item in desktop_risk_matrix
        )

    tgs_reference = normalized_form.get('tgs_reference', tgs_sheet_reference)
    tgs_number = normalized_form.get('tgs_number', '')
    cc_input = normalized_form.get('cc', normalized_form.get('cc_number', ''))
    revision_number = normalized_form.get('revision_number', '0')
    cc_match = re.search(r'CC\d{5}', cc_input, flags=re.IGNORECASE)
    cc_base_number = cc_match.group(0) if cc_match else cc_input
    document_reference = build_document_reference(cc_input or cc_base_number, revision_number)
    cc_number = document_reference or cc_base_number
    tmr_cert_number = normalized_form.get(
        'tmr_cert_82_number',
        normalized_form.get('tmr_cert82_number', nto_op_number or nto_tmr_tmd)
    )
    author_name = normalized_form.get('author_name', client_name or 'Crompton Concepts')
    author_signature = normalized_form.get('author_signature', initials_from_name(author_name))
    author_position = normalized_form.get('author_position', 'Traffic Technician')
    author_date = normalized_form.get('author_date', today)
    lot_parcel_number = normalized_form.get('lot_parcel_number', '')
    development_type = normalized_form.get('development_type', '')
    number_of_floors = normalized_form.get('number_of_floors', '')
    tgs_caption = ' '.join(part for part in [tgs_reference, tgs_number] if part).strip() or 'Traffic Guidance Scheme'
    general_ttm = normalized_form.get('traffic_management_notes', '') or (
        f'Traffic control measures will be implemented via {tgs_methodology} in accordance with the approved TGS and site conditions.'
        if tgs_methodology else
        'Traffic control measures will be implemented in accordance with the approved TGS and site conditions.'
    )
    stage_option_analysis = build_stage_option_analysis(stages, tgs_suite, tgs_methodology or construction_activity or general_ttm)
    stage_option_rows = [
        {
            'suite': tgs_suite or 'Civil Construction',
            'stage': stage,
            'methodology': tgs_methodology or construction_activity,
            'option_analysis': build_stage_option_analysis([stage], tgs_suite, tgs_methodology or construction_activity or general_ttm),
        }
        for stage in (stages or ([tgs_stage_name] if tgs_stage_name else ['General traffic management stage']))
    ]

    role_defaults = {
        'client_contact': client_name,
        'client_position': client_position,
        'client_phone': client_phone,
        'client_email': client_email,
        'pm_name': pm_name,
        'pm_position': pm_position,
        'pm_phone': pm_phone,
        'pm_email': pm_email,
        'sm_name': sm_name,
        'sm_position': sm_position,
        'sm_phone': sm_phone,
        'sm_email': sm_email,
        'se_name': se_name,
        'se_position': se_position,
        'se_phone': se_phone,
        'se_email': se_email,
        'wo_name': wo_name,
        'wo_position': wo_position,
        'wo_phone': wo_phone,
        'wo_email': wo_email,
        'nto_name': nto_name,
        'nto_position': nto_position,
        'nto_tmd_number': nto_tmd_number,
        'nto_op_number': nto_op_number,
        'nto_tmr_tmd': nto_tmr_tmd or nto_op_number or tmr_cert_number,
        'nto_phone': nto_phone,
        'nto_email': nto_email,
        'reviewer_name': reviewer_name,
        'reviewer_op': reviewer_op,
        'council_dept': council_dept,
        'council_officer': council_officer,
        'council_phone': council_phone,
        'council_email': council_email,
        'tc_provider_name': tc_provider_name,
        'tc_provider_phone': tc_provider_phone,
    }

    context = {field: '' for field in template_fields}
    context.update(normalized_form)

    defaults = {
        'project_name': project_name,
        'project_tile': project_name,
        'project_location': project_location,
        'project_location_e_g_12_example_road_sampletown': project_location,
        'street_address_suburb_state_postcode': project_location,
        'type_of_development_e_g_residential_commercial_mixed_use': development_type,
        'number_of_floors_e_g_3_storeys': number_of_floors,
        'insert_lot_and_parcel_number': lot_parcel_number,
        'client_company': client_company,
        'client_contracting_company_name': client_company,
        'tc_company': traffic_control_company,
        'cc': cc_number,
        'cc_number': cc_number,
        'cc_base_number': cc_base_number,
        'revision_number': revision_number,
        'revision_text': f'Rev {revision_number}' if revision_number else '',
        'document_reference': document_reference,
        'document_number': document_reference,
        'ctmp_reference': document_reference,
        'cc_revision_full': document_reference,
        'tmr_cert_82_number': tmr_cert_number,
        'tmr_tmd': tmr_cert_number,
        'author_name': author_name,
        'reviewer_name': reviewer_name,
        'reviewer_op': reviewer_op,
        'council_dept': council_dept,
        'council_officer': council_officer,
        'council_phone': council_phone,
        'council_email': council_email,
        'tc_provider_name': tc_provider_name or traffic_control_company,
        'tc_provider_phone': tc_provider_phone,
        'author_signature': author_signature,
        'author_position': author_position,
        'author_date': author_date,
        'construction_activity': construction_activity,
        'construction_suite': tgs_suite,
        'suite': tgs_suite,
        'tgs_stage_name': tgs_stage_name,
        'tgs_methodology': tgs_methodology,
        'tgs_site_location': tgs_site_location,
        'tgs_cross_streets': tgs_cross_streets,
        'tgs_road_speed': tgs_road_speed,
        'sheet_reference': tgs_sheet_reference,
        'methodology_implemented': tgs_methodology or construction_activity,
        'option_analysis': stage_option_analysis,
        'stage_option_analysis': stage_option_analysis,
        'staging_option_table': stage_option_rows,
        'stage_option_rows': stage_option_rows,
        'brief_description_of_the_project_including_nature_goals_and_any_major_components': project_description or scope_of_work or stage_text,
        'scope_of_work': scope_of_work or construction_activity or stage_text,
        'scope_of_works': scope_of_work or construction_activity or stage_text,
        'stages': stages,
        'enter_construction_stage': stage_text,
        'enter_construction_methodology_for_the_stage': tgs_methodology or construction_activity or stage_text,
        'enter_reasoning': stage_option_analysis,
        'enter_details': stage_option_analysis,
        'enter_details_reasoning': stage_option_analysis,
        'enter_activity_details': tgs_methodology or construction_activity or stage_option_analysis,
        'enter_proposed_control_method': tgs_methodology or general_ttm,
        'start_date': start_date,
        'end_date': end_date,
        'duration': duration,
        'date': today,
        'enter_date': today,
        'release_date': today,
        'current_year': current_year,
        'insert_current_year': current_year,
        'name': generic_name,
        'position': generic_position,
        'phone': generic_phone,
        'email': generic_email,
        'local_government_area': local_government_area,
        'describe_lga': local_government_area,
        'jurisdiction': jurisdiction,
        'target_year': target_year,
        'aadt_target_year': tia_metrics.get('aadt_target_year', ''),
        'vcr_score': tia_metrics.get('vcr_score', ''),
        'queue_length_m': tia_metrics.get('queue_length_m', ''),
        'los_rating': tia_metrics.get('los_rating', ''),
        'traffic_impact_summary': (
            f"Target year {target_year} AADT: {tia_metrics.get('aadt_target_year', '')}; "
            f"VCR: {tia_metrics.get('vcr_score', '')}; "
            f"Queue Length: {tia_metrics.get('queue_length_m', '')} m; "
            f"LOS: {tia_metrics.get('los_rating', '')}"
        ),
        'vehicular_options': structured_options.get('vehicular_options', []),
        'vulnerable_users': structured_options.get('vulnerable_users', []),
        'desktop_risk_assessment': risk_assessment_summary,
        'desktop_risk_matrix': desktop_risk_matrix,
        'describe_suburb': project_location,
        'suburb': project_location,
        'road': project_location,
        'road_name': project_location,
        'insert_location': project_location,
        'insert_road': project_location,
        'tgs_reference': tgs_reference,
        'tgs_number': tgs_number,
        'insert_caption_to_figure_by_right_clicking_the_image': tgs_caption if uploaded_images else '',
        'insert_ttm': general_ttm,
        'insert_active_ttm_measures': general_ttm,
        'detail_ttm_implemented_for_vehicular_traffic': general_ttm,
        'yes_no': 'No',
        'adopted_not_adopted': 'Not Adopted',
        'affected_unaffected': 'unaffected',
        'are_affected_or_unaffected': 'unaffected',
        'be_affected_be_unaffected': 'be unaffected',
        'delete_other_reviewers_as_required': '',
        'remove_if_not_provided': '',
        'remove_the_equipment_not_relevant_to_project': '',
        'required_for_gccc_remove_if_not_required': '',
    }

    if uploaded_images:
        defaults['insert_figure'] = InlineImage(doc, io.BytesIO(uploaded_images[0]['bytes']), width=Mm(150))

    defaults.update(role_defaults)
    context.update(defaults)

    structured_tables = {
        'document_preparation': parse_structured_rows(
            normalized_form.get('document_preparation'),
            5,
            default_rows=[
                ['Prepared By', author_name, author_signature, author_position, author_date],
                ['Reviewed By', reviewer_name, initials_from_name(reviewer_name), f'TMD ({reviewer_op})', author_date],
            ],
        ),
        'revision_history': parse_structured_rows(
            normalized_form.get('revision_history'),
            4,
            default_rows=[[revision_number, document_reference or project_name, 'TGS/TMP', author_date]],
        ),
        'distribution': parse_structured_rows(
            normalized_form.get('distribution_list'),
            5,
            default_rows=[[client_name, client_company, client_position, '1', author_date]],
        ),
        'change_log': parse_structured_rows(
            normalized_form.get('change_log'),
            3,
            default_rows=[['Initial Issue', 'All Sections', author_date]],
        ),
        'badge_contacts': parse_structured_rows(
            normalized_form.get('badge_contacts'),
            4,
            default_rows=[
                [normalized_form.get('pm_name', ''), 'Project Manager', normalized_form.get('pm_phone', ''), normalized_form.get('pm_email', '')],
                [normalized_form.get('sm_name', ''), 'Site Manager', normalized_form.get('sm_phone', ''), normalized_form.get('sm_email', '')],
                [normalized_form.get('wo_name', ''), 'Site Supervisor', normalized_form.get('wo_phone', ''), normalized_form.get('wo_email', '')],
                [normalized_form.get('client_contact', ''), 'WHS Officer', normalized_form.get('client_phone', ''), normalized_form.get('client_email', '')],
            ],
        ),
        'tm_consultants': parse_structured_rows(
            normalized_form.get('tm_consultants'),
            4,
            default_rows=[
                [author_name, author_position, '(07) 3187 8940', 'operations@cromptonconcepts.com.au'],
                ['Aaron Anthony', 'TMD (OP611)', '0413 232 019', 'aaron@cromptonconcepts.com.au'],
                ['Andrew Bryers', 'TMD (OP684)', '0450 920 096', 'andrew@cromptonconcepts.com.au'],
                ['Mark Bryers', 'TMD (OP449)', '0403 619 110', 'mark@cromptonconcepts.com.au'],
                ['Kelly Govind', 'TMD (OP442)', '0406 546 550', 'kelly@cromptonconcepts.com.au'],
                ['Sanju Bhandari', 'TMD (OP1569)', '0405 535 187', 'sanju@cromptonconcepts.com.au'],
            ],
        ),
        'authority_contacts': parse_structured_rows(
            normalized_form.get('authority_contacts'),
            4,
            default_rows=[[council_dept, council_officer, council_phone, council_email]],
        ),
        'dtmr_contacts': parse_structured_rows(
            normalized_form.get('dtmr_contacts'),
            4,
            default_rows=[['DTMR', 'Assessing Officer', normalized_form.get('dtmr_contact', ''), normalized_form.get('dtmr_email', '')]],
        ),
        'nto_contacts': parse_structured_rows(
            normalized_form.get('nto_contacts'),
            5,
            default_rows=[[nto_name, nto_position or 'TMD', nto_op_number or nto_tmd_number or nto_tmr_tmd or tmr_cert_number, nto_phone, nto_email]],
        ),
        'traffic_control_contacts': parse_structured_rows(
            normalized_form.get('traffic_control_contacts'),
            3,
            default_rows=[
                [tc_provider_name or normalized_form.get('sm_name', ''), 'Traffic Control Provider', tc_provider_phone or normalized_form.get('sm_phone', '')],
                [normalized_form.get('wo_name', normalized_form.get('pm_name', '')), 'Site Supervisor', normalized_form.get('wo_phone', normalized_form.get('pm_phone', ''))],
            ],
        ),
        'emergency_contacts': parse_structured_rows(
            normalized_form.get('emergency_contacts'),
            3,
            default_rows=[[client_name, client_position, client_phone]],
        ),
    }

    loop_context = build_loop_context(structured_tables)
    context['_structured_tables'] = structured_tables
    context.update(loop_context)

    for field in template_fields:
        if context.get(field):
            continue

        if field.startswith('e_g_') or field.startswith('remove_') or 'delete_other_reviewers' in field:
            context[field] = ''
        elif 'project' in field and ('name' in field or 'tile' in field):
            context[field] = project_name
        elif 'project' in field and 'location' in field:
            context[field] = project_location
        elif 'client' in field and 'company' in field:
            context[field] = client_company
        elif 'scope' in field or 'brief_description' in field:
            context[field] = scope_of_work or project_description or stage_text
        elif 'methodology' in field or 'proposed_control_method' in field:
            context[field] = tgs_methodology or construction_activity or general_ttm
        elif 'construction' in field or 'activity' in field or 'stage' in field:
            context[field] = construction_activity or stage_text
        elif field == 'date' or field.endswith('_date'):
            context[field] = today
        elif 'duration' in field:
            context[field] = duration
        elif 'lga' in field or 'local_government_area' in field:
            context[field] = local_government_area
        elif 'suburb' in field or 'location' in field or 'road' in field:
            context[field] = project_location
        elif 'tgs_reference' in field:
            context[field] = tgs_reference
        elif 'tgs_number' in field or field == 'number':
            context[field] = tgs_number
        elif field in role_defaults:
            context[field] = role_defaults[field]
        elif 'phone' in field:
            context[field] = generic_phone
        elif 'email' in field:
            context[field] = generic_email
        elif 'position' in field:
            context[field] = generic_position
        elif field.endswith('_name') or field == 'name':
            context[field] = generic_name or project_name
        elif 'yes_no' in field:
            context[field] = 'No'
        elif 'adopted' in field:
            context[field] = 'Not Adopted'
        elif 'affected' in field:
            context[field] = 'unaffected'
        elif 'ttm' in field or 'traffic_control' in field or 'control_measure' in field:
            context[field] = general_ttm
        elif 'reasoning' in field or 'details' in field or 'option_analysis' in field:
            context[field] = stage_option_analysis

    return context


def cleanup_rendered_document(document, context, has_uploaded_images):
    delete_contains = (
        'replicate this whole section',
        'remove if not provided',
        'required for gccc',
        'remove the equipment not relevant to project',
        'delete otherwise',
    )
    delete_exact = {
        'dd/mm/yy',
        'date',
        ':',
        '-',
    }

    for paragraph in list(iter_paragraphs(document)):
        original_text = paragraph.text
        updated_text = replace_literal_placeholders(original_text, context)
        if updated_text != original_text:
            paragraph.text = updated_text

        for run in paragraph.runs:
            if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                run.font.highlight_color = None

        text = ' '.join(paragraph.text.split())
        if not text:
            continue

        lower = text.lower()
        if any(marker in lower for marker in delete_contains):
            delete_paragraph(paragraph)
            continue

        has_placeholder_artifact = ('<<' in original_text or '>>' in original_text or '{{' in original_text or '}}' in original_text)
        if has_placeholder_artifact and lower in delete_exact:
            delete_paragraph(paragraph)
            continue

        if not has_uploaded_images and has_placeholder_artifact and ('site active' in lower or 'site inactive' in lower or lower == 'plan'):
            delete_paragraph(paragraph)

    enable_update_fields(document)


def extract_tgs_page_images(request_files, render_scale=2):
    page_images = []
    pdf_file = request_files.get('tgsPdf')

    if pdf_file and pdf_file.filename:
        pdf_bytes = pdf_file.read()
        if pdf_bytes:
            pdf_doc = fitz.open(stream=pdf_bytes, filetype='pdf')
            try:
                tgs_pattern = re.compile(r'CC\d{5}-S\d+-?\d+[A-C]?', re.IGNORECASE)

                for page_index in range(pdf_doc.page_count):
                    page = pdf_doc.load_page(page_index)
                    raw_text = page.get_text('text')
                    page_text = ' '.join(raw_text.split())
                    upper_text = page_text.upper()
                    has_tgs_sheet = bool(tgs_pattern.search(upper_text))
                    page_metadata = extract_tgs_metadata_from_text(raw_text)
                    has_title_block_metadata = bool(
                        page_metadata.get('sheet_reference') or
                        page_metadata.get('site_location') or
                        page_metadata.get('cross_streets') or
                        page_metadata.get('road_speed') or
                        page_metadata.get('stage_name') or
                        page_metadata.get('methodology')
                    )

                    image_count = len(page.get_images(full=True))
                    drawing_count = len(page.get_drawings())
                    low_text_density = len(page_text) < 1200
                    drawing_page = (image_count + drawing_count) > 0 and low_text_density
                    has_traffic_keywords = any(keyword in upper_text for keyword in (
                        'ROAD SPEED', 'DETOUR', 'STOP/SLOW', 'STOP SLOW', 'PEDESTRIAN MANAGEMENT',
                        'TRAFFIC MANAGEMENT', 'LANE CLOSURE', 'FOOTPATH WORKS', 'TGS'
                    ))
                    looks_like_contents_page = ('CONTENTS' in upper_text) or (upper_text.count('FIGURE ') >= 3) or bool(re.search(r'\.{4,}', page_text))

                    if page_index == 0 and not has_tgs_sheet and not has_title_block_metadata:
                        continue

                    if looks_like_contents_page and not has_tgs_sheet and not has_title_block_metadata:
                        continue

                    if not has_tgs_sheet and not has_title_block_metadata and not (drawing_page and has_traffic_keywords):
                        continue

                    pixmap = page.get_pixmap(matrix=fitz.Matrix(render_scale, render_scale), alpha=False)
                    page_images.append({
                        'name': f'{pdf_file.filename}-page-{page_index + 1}.png',
                        'bytes': pixmap.tobytes('png'),
                        'page_number': page_index + 1,
                        'page_text': page_text,
                        'suite': page_metadata.get('suite', ''),
                        'stage_name': page_metadata.get('stage_name', ''),
                        'methodology': page_metadata.get('methodology', ''),
                        'site_location': page_metadata.get('site_location', ''),
                        'cross_streets': page_metadata.get('cross_streets', ''),
                        'road_speed': page_metadata.get('road_speed', ''),
                        'sheet_reference': page_metadata.get('sheet_reference', ''),
                    })
            finally:
                pdf_doc.close()

    if not page_images:
        for file in request_files.getlist('tgsDiagrams'):
            if file and file.filename:
                file_bytes = file.read()
                if file_bytes:
                    page_images.append({'name': file.filename, 'bytes': file_bytes, 'page_number': len(page_images) + 1})

    return page_images


def append_tgs_appendix(document, uploaded_images, caption_prefix):
    if not uploaded_images:
        return

    document.add_page_break()
    document.add_heading('Traffic Guidance Schemes', level=1)

    for index, image_info in enumerate(uploaded_images, start=1):
        source_page = image_info.get('page_number', index)
        label = caption_prefix if len(uploaded_images) == 1 else f'{caption_prefix} - Page {source_page}'
        document.add_paragraph(label)
        document.add_picture(io.BytesIO(image_info['bytes']), width=Mm(160))


@app.route('/')
def serve_html():
    return send_file(HTML_FILE)


@app.route('/api/template-index')
def get_template_index():
    template_name = request.args.get('template')

    if template_name:
        template_name = resolve_template_name(template_name)
        template_path = os.path.join(TEMPLATE_DIR, template_name)
        if not os.path.exists(template_path):
            return jsonify({'error': f'Template {template_name} not found'}), 404
        return jsonify(build_template_index(template_name))

    template_index = {}
    for filename in sorted(os.listdir(TEMPLATE_DIR)):
        if filename.lower().endswith('.docx'):
            template_index[filename] = build_template_index(filename)

    return jsonify(template_index)


@app.route('/api/analyze-tgs', methods=['POST'])
def analyze_tgs():
    try:
        uploaded_images = extract_tgs_page_images(request.files)
        if not uploaded_images:
            return jsonify({'error': 'No TGS PDF or diagram pages were found'}), 400
        return jsonify(summarize_tgs_analysis(uploaded_images))
    except Exception as error:
        print(f'Error analysing TGS: {error}')
        message = str(error) if app.config['DEBUG'] else 'TGS analysis failed. Check server logs.'
        return jsonify({'error': message}), 500


def build_tgs_preview_payload(uploaded_images):
    preview_pages = []

    for index, image_info in enumerate(uploaded_images or [], start=1):
        page_number = image_info.get('page_number', index)
        sheet_reference = normalize_context_value(image_info.get('sheet_reference', '')).strip()
        stage_name = normalize_context_value(image_info.get('stage_name', '')).strip()
        methodology = normalize_context_value(image_info.get('methodology', '')).strip()
        subtitle_parts = [part for part in [stage_name, methodology] if part]

        preview_pages.append({
            'id': f'tgs-{index}',
            'pageNumber': page_number,
            'sheetReference': sheet_reference,
            'stageName': stage_name,
            'methodology': methodology,
            'label': sheet_reference or f'TGS Page {page_number}',
            'subtitle': ' • '.join(subtitle_parts),
            'imageUrl': f"data:image/png;base64,{base64.b64encode(image_info['bytes']).decode('ascii')}",
        })

    return preview_pages


@app.route('/api/tgs-preview', methods=['POST'])
def tgs_preview():
    try:
        uploaded_images = extract_tgs_page_images(request.files, render_scale=1.15)
        if not uploaded_images:
            return jsonify({'error': 'No TGS PDF or diagram pages were found'}), 400

        response = summarize_tgs_analysis(uploaded_images)
        response['pages'] = build_tgs_preview_payload(uploaded_images)
        return jsonify(response)
    except Exception as error:
        print(f'Error building TGS preview: {error}')
        message = str(error) if app.config['DEBUG'] else 'TGS preview failed. Check server logs.'
        return jsonify({'error': message}), 500


@app.route('/api/generate-ctmp', methods=['POST'])
@app.route('/generate', methods=['POST'])
def generate_ctmp():
    try:
        template_name = request.form.get('templateType')
        project_name = request.form.get('projectName', 'Project')
        data_source = request.form.get('dataSource', 'hourly_profile')
        stages = request.form.getlist('stages')

        if not template_name:
            return jsonify({'error': 'No template selected'}), 400

        template_name = resolve_template_name(template_name)
        template_path = os.path.join(TEMPLATE_DIR, template_name)
        if not os.path.exists(template_path):
            return jsonify({'error': f'Template {template_name} not found'}), 404

        doc = DocxTemplate(template_path)

        uploaded_images = extract_tgs_page_images(request.files)

        template_fields = extract_template_fields(template_path)
        context = build_context(request.form, template_fields, stages, doc, uploaded_images)

        jurisdiction = request.form.get('jurisdiction', 'TMR')
        target_year = request.form.get('targetYear', request.form.get('target_year', '2026'))
        hourly_db = request.form.get('hourlyDbData', request.form.get('hourly_database_data'))
        calculated_est = request.form.get('calculatedEstimates', request.form.get('calculated_estimates'))

        tia_results = calculate_traffic_impact(jurisdiction, hourly_db, calculated_est, target_year)
        options_analysis = generate_options_analysis(context.get('stages', []))
        risk_matrix = generate_risk_matrix()

        context.update({
            'target_year': target_year,
            'jurisdiction': jurisdiction,
            'tia_target_year': target_year,
            'tia_aadt': tia_results.get('aadt_target_year'),
            'tia_vcr': tia_results.get('vcr_score'),
            'tia_queue': tia_results.get('queue_length_m'),
            'tia_queue_length': tia_results.get('queue_length_m'),
            'tia_los': tia_results.get('los_rating'),
            'vehicular_options': options_analysis.get('vehicular_options'),
            'vulnerable_options': options_analysis.get('vulnerable_users'),
            'pedestrian_options': options_analysis.get('pedestrian'),
            'risk_matrix': risk_matrix,
            'risk_assessment_table': risk_matrix,
        })

        context['data_source'] = 'Hourly Profile Target Year' if data_source == 'hourly_profile' else 'Standard Estimate'
        context['tgs_images'] = [InlineImage(doc, io.BytesIO(image['bytes']), width=Mm(150)) for image in uploaded_images]

        doc.render(context)

        rendered_stream = io.BytesIO()
        doc.save(rendered_stream)
        rendered_stream.seek(0)

        final_doc = Document(rendered_stream)
        populate_structured_tables(final_doc, context)
        cleanup_rendered_document(final_doc, context, has_uploaded_images=bool(uploaded_images))
        append_tgs_appendix(final_doc, uploaded_images, context.get('document_reference') or context.get('tgs_reference') or 'Traffic Guidance Scheme')

        file_stream = io.BytesIO()
        final_doc.save(file_stream)
        file_stream.seek(0)

        safe_filename = f"{project_name.replace(' ', '_')}_CTMP.docx"
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=safe_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as error:
        print(f'Error generating document: {error}')
        message = str(error) if app.config['DEBUG'] else 'Document generation failed. Check server logs.'
        return jsonify({'error': message}), 500


if __name__ == '__main__':
    app.run(debug=app.config['DEBUG'], port=5000)